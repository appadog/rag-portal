"""Real document extraction with an explicit, observable OCR fallback."""

from __future__ import annotations

import base64
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path


class DocumentParseError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    parser: str
    used_ocr: bool = False
    warnings: list[str] = field(default_factory=list)


def decode_source(content: str | None, content_base64: str | None) -> bytes:
    if content_base64:
        try:
            return base64.b64decode(content_base64, validate=True)
        except ValueError as error:
            raise DocumentParseError("파일 데이터가 손상되어 읽을 수 없습니다.") from error
    return (content or "").encode("utf-8")


def extract_document(
    *, filename: str, content_type: str, content: str | None, content_base64: str | None
) -> ParsedDocument:
    data = decode_source(content, content_base64)
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf" or content_type == "application/pdf":
        return extract_pdf(data)
    if suffix == ".docx" or "wordprocessingml" in content_type:
        return extract_docx(data)
    if suffix in {".xlsx", ".xlsm"} or "spreadsheetml" in content_type:
        return extract_xlsx(data)
    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tiff", ".bmp"} or content_type.startswith("image/"):
        return extract_image_ocr(data)
    text = data.decode("utf-8", errors="replace").strip()
    if not text:
        raise DocumentParseError("텍스트를 추출하지 못했습니다. 빈 파일인지 확인해 주세요.")
    return ParsedDocument(text=text, parser="text")


def extract_pdf(data: bytes) -> ParsedDocument:
    from pypdf import PdfReader

    try:
        reader = PdfReader(BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as error:
        raise DocumentParseError("PDF를 읽지 못했습니다. 암호화 또는 손상 여부를 확인해 주세요.") from error
    text = "\n\n".join(page for page in pages if page)
    if text:
        return ParsedDocument(text=text, parser="pypdf")
    return extract_pdf_ocr(data)


def extract_pdf_ocr(data: bytes) -> ParsedDocument:
    try:
        import fitz
        from PIL import Image
        import pytesseract

        document = fitz.open(stream=data, filetype="pdf")
        pages: list[str] = []
        for page in document:
            pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
            image = Image.open(BytesIO(pixmap.tobytes("png")))
            pages.append(pytesseract.image_to_string(image, lang="kor+eng").strip())
    except Exception as error:
        raise DocumentParseError(
            "스캔 PDF입니다. OCR 실행 환경(Tesseract 언어팩 kor+eng)을 설치한 뒤 다시 시도해 주세요."
        ) from error
    text = "\n\n".join(page for page in pages if page)
    if not text:
        raise DocumentParseError("OCR을 실행했지만 읽을 수 있는 텍스트를 찾지 못했습니다.")
    return ParsedDocument(text=text, parser="pymupdf+tesseract", used_ocr=True)


def extract_docx(data: bytes) -> ParsedDocument:
    from docx import Document

    try:
        document = Document(BytesIO(data))
    except Exception as error:
        raise DocumentParseError("DOCX를 읽지 못했습니다. 파일 형식을 확인해 주세요.") from error
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        blocks.extend(row for row in rows if row.strip())
    text = "\n\n".join(blocks)
    if not text:
        raise DocumentParseError("DOCX에서 읽을 수 있는 본문이나 표를 찾지 못했습니다.")
    return ParsedDocument(text=text, parser="python-docx")


def extract_xlsx(data: bytes) -> ParsedDocument:
    from openpyxl import load_workbook

    try:
        workbook = load_workbook(BytesIO(data), read_only=True, data_only=True)
    except Exception as error:
        raise DocumentParseError("XLSX를 읽지 못했습니다. 파일 형식을 확인해 주세요.") from error
    sheets: list[str] = []
    for worksheet in workbook.worksheets:
        rows = []
        for row in worksheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                rows.append(" | ".join(values))
        if rows:
            sheets.append(f"[시트: {worksheet.title}]\n" + "\n".join(rows))
    text = "\n\n".join(sheets)
    if not text:
        raise DocumentParseError("XLSX에서 읽을 수 있는 셀 값을 찾지 못했습니다.")
    return ParsedDocument(text=text, parser="openpyxl")


def extract_image_ocr(data: bytes) -> ParsedDocument:
    try:
        from PIL import Image
        import pytesseract

        text = pytesseract.image_to_string(Image.open(BytesIO(data)), lang="kor+eng").strip()
    except Exception as error:
        raise DocumentParseError("이미지 OCR을 실행하려면 Tesseract 언어팩 kor+eng이 필요합니다.") from error
    if not text:
        raise DocumentParseError("OCR을 실행했지만 읽을 수 있는 텍스트를 찾지 못했습니다.")
    return ParsedDocument(text=text, parser="pillow+tesseract", used_ocr=True)
