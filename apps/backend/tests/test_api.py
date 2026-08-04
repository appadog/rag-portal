import base64
from io import BytesIO
import time

import pytest
from fastapi.testclient import TestClient
from docx import Document as DocxDocument
import fitz
from openpyxl import Workbook

from app import main
from app.generation import GenerationEndpointError, GenerationResult, GenerationSentence
from app.main import app
from app.retrieval import bm25_scores, embed, rank
from app.source_storage import source_storage


client = TestClient(app)


def test_openapi_includes_workspace_state_contract() -> None:
    schema = client.get("/api/v1/openapi.json")
    assert schema.status_code == 200
    paths = schema.json()["paths"]
    assert "/api/v1/rag-instances/{instance_id}/jobs" in paths
    assert "/api/v1/rag-instances/{instance_id}/artifacts" in paths
    assert "/api/v1/documents/{document_id}/segments/{segment_id}" in paths
    assert "/api/v1/model-runtime" in paths
    assert "/api/v1/rag-instances/{instance_id}/execution-plan" in paths
    assert "/api/v1/large-document-policy" in paths
    assert "/api/v1/rag-instances/{instance_id}/documents/{document_id}/reparse" in paths
    assert "/api/v1/rag-instances/{instance_id}/retuning-recommendation" in paths
    assert "/api/v1/job-platform" in paths
    assert "/api/v1/rag-jobs/dead-letters" in paths
    assert "/api/v1/rag-jobs/{job_id}/recover" in paths
    assert "/api/v1/rag-instances/{instance_id}/candidate-exploration" in paths
    assert "/api/v1/candidate-exploration/{exploration_id}/rollback" in paths
    assert "/api/v1/candidate-exploration/{exploration_id}/restore" in paths
    assert "/api/v1/rag-instances/{instance_id}/search/preflight" in paths


def test_large_document_policy_defaults_to_500_chunks(monkeypatch) -> None:
    monkeypatch.delenv("RAG_PORTAL_COMPARISON_CHUNK_THRESHOLD", raising=False)
    policy = client.get("/api/v1/large-document-policy")
    assert policy.status_code == 200
    assert policy.json()["comparison_chunk_threshold"] == 500


def test_loopback_vite_ports_are_allowed_by_cors() -> None:
    response = client.options(
        "/api/v1/rag-instances",
        headers={
            "Origin": "http://127.0.0.1:5175",
            "Access-Control-Request-Method": "POST",
        },
    )

    assert response.status_code == 200
    assert response.headers["access-control-allow-origin"] == "http://127.0.0.1:5175"


def test_embedding_candidates_are_ranked_and_the_selected_candidate_is_persisted() -> None:
    recommendations = client.post(
        "/api/v1/rag-instances/embedding-recommendations",
        json={"primary_language": "ko", "requires_on_premise": True, "budget": "standard"},
    )
    assert recommendations.status_code == 200
    items = recommendations.json()["items"]
    assert [item["id"] for item in items] == [
        "Qwen3-Embedding-0.6B",
        "BGE-M3",
        "EmbeddingGemma-300M",
    ]
    assert items[0]["recommended"] is True
    assert all({"label", "reason", "tradeoff"} <= item.keys() for item in items)

    created = client.post(
        "/api/v1/rag-instances",
        json={
            "name": "후보 선택 RAG",
            "questionnaire": {
                "primary_language": "ko",
                "requires_on_premise": True,
                "embedding_model": "BGE-M3",
            },
        },
    )
    assert created.status_code == 201
    assert created.json()["embedding_model"] == "BGE-M3"
    assert len(created.json()["recommendation"]["candidates"]) == 3


def test_model_runtime_and_instance_execution_plan_explain_required_services() -> None:
    runtime = client.get("/api/v1/model-runtime")
    assert runtime.status_code == 200
    services = {service["key"]: service for service in runtime.json()["services"]}
    assert services["embedding-bge-m3"]["technique"] == "embedding"
    assert services["reranker-bge-m3"]["technique"] == "reranking"
    assert services["ocr-tesseract-kor-eng"]["technique"] == "ocr"

    instance_id = create_instance()
    plan = client.get(f"/api/v1/rag-instances/{instance_id}/execution-plan")
    assert plan.status_code == 200
    keys = {service["key"] for service in plan.json()["required_services"]}
    assert {"embedding-bge-m3", "reranker-bge-m3", "generator-grounded"} <= keys


def test_embedding_benchmark_persists_model_quality_contract() -> None:
    created = client.post("/api/v1/embedding-benchmarks/run")
    assert created.status_code == 200
    body = created.json()
    assert body["run"]["query_count"] == 3
    assert {result["model_id"] for result in body["results"]} == {
        "BGE-M3",
        "Qwen3-Embedding-0.6B",
        "EmbeddingGemma-300M",
    }
    assert all("recall_at_5" in result and "provider" in result for result in body["results"])
    latest = client.get("/api/v1/embedding-benchmarks/latest")
    assert latest.status_code == 200
    assert latest.json()["run"]["id"] == body["run"]["id"]


def test_docx_and_xlsx_are_extracted_before_candidate_indexing() -> None:
    docx = DocxDocument()
    docx.add_heading("출장 규정", level=1)
    docx.add_paragraph("국내 출장 숙박비는 1박 10만원을 한도로 합니다.")
    docx_table = docx.add_table(rows=2, cols=2)
    docx_table.cell(0, 0).text = "항목"
    docx_table.cell(0, 1).text = "한도"
    docx_table.cell(1, 0).text = "식비"
    docx_table.cell(1, 1).text = "150달러"
    docx_buffer = BytesIO()
    docx.save(docx_buffer)

    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "국내 출장"
    worksheet.append(["항목", "한도"])
    worksheet.append(["숙박비", "10만원"])
    xlsx_buffer = BytesIO()
    workbook.save(xlsx_buffer)

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Domestic travel lodging is limited to 100,000 won.")
    pdf_bytes = pdf.tobytes()

    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={
            "documents": [
                {
                    "filename": "travel.docx",
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "content_base64": base64.b64encode(docx_buffer.getvalue()).decode(),
                },
                {
                    "filename": "travel.xlsx",
                    "content_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    "content_base64": base64.b64encode(xlsx_buffer.getvalue()).decode(),
                },
                {
                    "filename": "travel.pdf",
                    "content_type": "application/pdf",
                    "content_base64": base64.b64encode(pdf_bytes).decode(),
                },
            ]
        },
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    documents = main.INSTANCES[instance_id].documents
    parsed_docx = next(document for document in documents.values() if document.filename.endswith(".docx"))
    parsed_xlsx = next(document for document in documents.values() if document.filename.endswith(".xlsx"))
    parsed_pdf = next(document for document in documents.values() if document.filename.endswith(".pdf"))
    assert parsed_docx.parser == "python-docx"
    assert "150달러" in parsed_docx.content
    assert parsed_xlsx.parser == "openpyxl"
    assert "숙박비 | 10만원" in parsed_xlsx.content
    assert parsed_pdf.parser == "pypdf"
    assert "lodging" in parsed_pdf.content
    assert all(main.CANDIDATES[candidate_id].vectors for document in documents.values() for candidate_id in document.candidate_ids)


def test_processing_job_can_be_cancelled_and_retried() -> None:
    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "retry.txt", "content": "국내 출장 숙박비는 1박 10만원입니다."}]},
    ).json()
    job_id = upload["job"]["id"]
    assert wait_for_job(job_id)["state"] == "SUCCEEDED"
    job = main.JOBS[job_id]

    job.state = main.JobState.QUEUED
    cancelled = client.post(f"/api/v1/rag-jobs/{job_id}/cancel")
    assert cancelled.status_code == 200
    assert cancelled.json()["can_cancel"] is True
    main.run_processing_job(job_id, job.pipeline_mode, job.reuse_source_document_id)
    assert client.get(f"/api/v1/rag-jobs/{job_id}").json()["state"] == "CANCELLED"

    retried = client.post(f"/api/v1/rag-jobs/{job_id}/retry")
    assert retried.status_code == 200
    assert retried.json()["attempt"] == 2
    assert wait_for_job(job_id)["state"] == "SUCCEEDED"


def test_operational_job_contract_dead_letters_and_explicit_recovery(monkeypatch) -> None:
    instance_id = create_instance()
    job = main.ProcessingJob(
        id=main.new_id(),
        instance_id=instance_id,
        document_ids=[],
        state=main.JobState.FAILED,
        current_step="실패",
        completed_units=0,
        total_units=3,
        created_at=main.datetime.now(main.UTC),
        max_attempts=2,
        attempt=2,
    )
    main.ensure_job_operational_fields(job)
    main.JOBS[job.id] = job
    main.mark_job_failure(job, "테스트 실패")
    assert job.state == main.JobState.DEAD_LETTER

    payload = client.get(f"/api/v1/rag-jobs/{job.id}")
    assert payload.status_code == 200
    assert payload.json()["can_recover"] is True
    assert payload.json()["operational"]["dead_letter_reason"] == "MAX_ATTEMPTS_EXHAUSTED"

    dead_letters = client.get("/api/v1/rag-jobs/dead-letters")
    assert dead_letters.status_code == 200
    assert any(item["id"] == job.id for item in dead_letters.json()["items"])

    dispatched: list[str] = []
    monkeypatch.setattr(main, "start_processing_job", lambda recovered, *_: dispatched.append(recovered.id))
    recovered = client.post(f"/api/v1/rag-jobs/{job.id}/recover")
    assert recovered.status_code == 200
    assert recovered.json()["state"] == "QUEUED"
    assert recovered.json()["attempt"] == 1
    assert dispatched == [job.id]


def test_job_retry_backoff_and_cancellation_checkpoint() -> None:
    instance_id = create_instance()
    job = main.ProcessingJob(
        id=main.new_id(),
        instance_id=instance_id,
        document_ids=[],
        state=main.JobState.FAILED,
        current_step="실패",
        completed_units=0,
        total_units=3,
        created_at=main.datetime.now(main.UTC),
        max_attempts=3,
        retry_backoff_seconds=60,
    )
    main.ensure_job_operational_fields(job)
    main.mark_job_failure(job, "일시 오류")
    main.JOBS[job.id] = job
    blocked = client.post(f"/api/v1/rag-jobs/{job.id}/retry")
    assert blocked.status_code == 409
    assert blocked.json()["detail"]["code"] == "JOB_RETRY_BACKOFF"

    job.cancel_requested = True
    with pytest.raises(main.JobCancelled):
        main.ensure_job_is_active(job)
    assert job.last_heartbeat_at is not None


def test_job_platform_reports_local_adapter_fallback(monkeypatch) -> None:
    monkeypatch.setenv("RAG_QUEUE_BACKEND", "unsupported-adapter")
    response = client.get("/api/v1/job-platform")
    assert response.status_code == 200
    body = response.json()
    assert body["queue"]["effective"] == "thread"
    assert body["queue"]["ready"] is False


def test_job_operational_environment_defaults_tolerate_invalid_values(monkeypatch) -> None:
    monkeypatch.setenv("RAG_JOB_MAX_ATTEMPTS", "invalid")
    monkeypatch.setenv("RAG_JOB_RETRY_BACKOFF_SECONDS", "invalid")
    assert main.job_max_attempts() == 3
    assert main.job_retry_backoff_seconds() == 0


def test_adaptive_candidate_exploration_never_auto_selects_and_can_roll_back_restore() -> None:
    instance_id = create_instance()
    uploaded = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "explore.txt", "content": "제1조 국내 출장 숙박비는 1박 10만원입니다. 제2조 식비는 1일 150달러입니다."}]},
    ).json()
    assert wait_for_job(uploaded["job"]["id"])["state"] == "SUCCEEDED"
    document_id = uploaded["documents"][0]["id"]
    document = main.INSTANCES[instance_id].documents[document_id]
    initial_candidate_ids = list(document.candidate_ids)
    initial_votes = {candidate_id: main.CANDIDATES[candidate_id].selection_count for candidate_id in initial_candidate_ids}

    created = client.post(
        f"/api/v1/rag-instances/{instance_id}/candidate-exploration",
        json={"document_ids": [document_id], "question": "숙박비 한도는?", "max_proposals": 2},
    )
    assert created.status_code == 201
    exploration = created.json()["candidate_exploration"]
    assert exploration["selection"]["automatic"] is False
    assert exploration["status"] == "PROPOSED"
    assert len(exploration["pool"]) == len(initial_candidate_ids)
    assert len(exploration["proposed"]) == 2
    assert document.finalized_candidate_id is None
    assert {candidate_id: main.CANDIDATES[candidate_id].selection_count for candidate_id in initial_candidate_ids} == initial_votes
    assert all(item["candidate"]["exploration"]["round_id"] == exploration["id"] for item in exploration["proposed"])
    provenance = client.get(f"/api/v1/rag-instances/{instance_id}").json()["documents"][0]["provenance"]["model"]
    assert {item["candidate"]["retrieval_config"] for item in exploration["proposed"]} <= set(provenance["retrieval_configs"])

    listed = client.get(f"/api/v1/rag-instances/{instance_id}/candidate-exploration")
    assert listed.status_code == 200
    assert listed.json()["items"][0]["id"] == exploration["id"]

    rolled_back = client.post(f"/api/v1/candidate-exploration/{exploration['id']}/rollback")
    assert rolled_back.status_code == 200
    assert rolled_back.json()["candidate_exploration"]["status"] == "ROLLED_BACK"
    assert all(main.CANDIDATES[candidate_id].archived for candidate_id in rolled_back.json()["archived_candidate_ids"])
    assert all(candidate_id in document.candidate_ids for candidate_id in initial_candidate_ids)

    restored = client.post(f"/api/v1/candidate-exploration/{exploration['id']}/restore")
    assert restored.status_code == 200
    assert restored.json()["candidate_exploration"]["status"] == "RESTORED"
    assert all(not main.CANDIDATES[candidate_id].archived for candidate_id in restored.json()["restored_candidate_ids"])


def test_exploration_preserves_large_document_sample_bound(monkeypatch) -> None:
    monkeypatch.setenv("RAG_PORTAL_COMPARISON_CHUNK_THRESHOLD", "3")
    instance_id = create_instance()
    content = "\n\n".join(f"제{number}조 국내 출장 숙박비는 1박 10만원입니다." for number in range(1, 14))
    uploaded = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "large-explore.txt", "content": content}]},
    ).json()
    assert wait_for_job(uploaded["job"]["id"])["state"] == "SUCCEEDED"
    document_id = uploaded["documents"][0]["id"]
    exploration = client.post(
        f"/api/v1/rag-instances/{instance_id}/candidate-exploration",
        json={"document_ids": [document_id], "max_proposals": 2},
    )
    assert exploration.status_code == 201
    proposed = exploration.json()["candidate_exploration"]["proposed"]
    assert all(item["candidate"]["comparison"]["scope"] == "SAMPLE" for item in proposed)
    assert all(item["candidate"]["chunk_count"] <= 3 for item in proposed)
    assert all(item["candidate"]["comparison"]["estimated_chunk_count"] >= item["candidate"]["chunk_count"] for item in proposed)


def test_search_preflight_reports_per_document_conflicts_without_generating(monkeypatch) -> None:
    monkeypatch.setenv("RAG_PORTAL_COMPARISON_CHUNK_THRESHOLD", "3")
    instance_id = create_instance()
    content = "\n\n".join(f"제{number}조 국내 출장 숙박비는 1박 10만원입니다." for number in range(1, 14))
    uploaded = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "preflight-large.txt", "content": content}]},
    ).json()
    assert wait_for_job(uploaded["job"]["id"])["state"] == "SUCCEEDED"
    document_id = uploaded["documents"][0]["id"]
    compared = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": "숙박비는?"},
    ).json()
    winner = next(item for item in compared["results"] if item["candidate_state"] == "READY")
    assert client.post(f"/api/v1/tuning-rounds/{compared['round']['id']}/vote", json={"candidate_ids": [winner["candidate"]["id"]]}).status_code == 200
    monkeypatch.setattr(main, "start_full_reindex_job", lambda *_: None)
    finalized = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/finalize", json={"document_id": document_id})
    assert finalized.status_code == 200
    before_artifacts = len(main.INSTANCES[instance_id].artifact_ids)

    preflight = client.get(
        f"/api/v1/rag-instances/{instance_id}/search/preflight",
        params=[("document_ids", document_id), ("document_ids", "missing-document")],
    )
    assert preflight.status_code == 200
    body = preflight.json()
    assert body["eligible"] is False
    conflicts = {item["document_id"]: item for item in body["conflicts"]}
    assert conflicts[document_id]["code"] == "FULL_REINDEX_PENDING"
    assert conflicts[document_id]["action"] == "WAIT_FOR_FULL_REINDEX"
    assert conflicts["missing-document"]["code"] == "DOCUMENT_NOT_FOUND"
    assert conflicts["missing-document"]["action"] == "REMOVE_DOCUMENT"
    assert len(main.INSTANCES[instance_id].artifact_ids) == before_artifacts

    blocked_search = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [document_id], "question": "숙박비는?"},
    )
    assert blocked_search.status_code == 409
    assert blocked_search.json()["detail"]["code"] == "FULL_REINDEX_PENDING"


def test_bm25_dense_hybrid_and_rerank_use_the_prepared_vector_index() -> None:
    texts = ["국내 출장 숙박비는 1박 10만원입니다.", "해외 출장 식비는 150달러입니다."]
    vectors = embed(texts, "BGE-M3").vectors
    assert bm25_scores("국내 출장 숙박비", texts)[0] > bm25_scores("국내 출장 숙박비", texts)[1]
    for retrieval_config in ("bm25", "dense", "hybrid", "hybrid_rerank"):
        scores, query_batch, metadata = rank(
            query="국내 출장 숙박비",
            texts=texts,
            vectors=vectors,
            retrieval_config=retrieval_config,
            model_name="BGE-M3",
        )
        assert len(scores) == 2
        assert query_batch.dimension == len(vectors[0])
        assert "reranker_provider" in metadata


def create_instance() -> str:
    response = client.post("/api/v1/rag-instances", json={"name": "출장 규정", "questionnaire": {"primary_language": "ko", "multi_hop_questions": False}})
    assert response.status_code == 201
    assert response.json()["embedding_model"] == "BGE-M3"
    return response.json()["id"]


def wait_for_job(job_id: str) -> dict:
    for _ in range(50):
        response = client.get(f"/api/v1/rag-jobs/{job_id}")
        assert response.status_code == 200
        job = response.json()
        if job["state"] in {"SUCCEEDED", "FAILED"}:
            return job
        time.sleep(0.01)
    raise AssertionError("job did not reach a terminal state")


def test_create_upload_compare_vote_finalize_and_search() -> None:
    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "출장비_규정.txt", "content": "제5조 해외 출장 식비는 국가 등급에 따라 1일 80~150달러를 한도로 지급합니다.\n\n제7조 국내 출장 숙박비는 1박 10만원을 한도로 합니다."}]},
    )
    assert upload.status_code == 202
    assert upload.json()["job"]["state"] in {"QUEUED", "PARSING", "GENERATING_CANDIDATES", "INDEXING"}
    assert wait_for_job(upload.json()["job"]["id"])["state"] == "SUCCEEDED"
    document_id = upload.json()["documents"][0]["id"]

    comparison = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/compare", json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"})
    assert comparison.status_code == 200
    first = comparison.json()["results"][0]
    assert "candidate" in first and "citations" in first

    vote = client.post(f"/api/v1/tuning-rounds/{comparison.json()['round']['id']}/vote", json={"candidate_ids": [first["candidate"]["id"]]})
    assert vote.status_code == 200
    assert vote.json()["tuning_status"]["can_finalize"] is True

    finalized = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/finalize", json={"document_id": document_id})
    assert finalized.status_code == 200
    assert finalized.json()["instance"]["status"] == "READY"

    searched = client.post(f"/api/v1/rag-instances/{instance_id}/search", json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?", "sensitivity": "balanced"})
    assert searched.status_code == 200
    assert searched.json()["grounded"] is True
    assert "10만원" in searched.json()["answer"]


def test_candidate_readiness_and_no_evidence_state_are_exposed_and_not_selectable() -> None:
    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "travel.txt", "content": "국내 출장 숙박비는 1박 10만원입니다."}]},
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    document_id = upload["documents"][0]["id"]

    detail = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    candidates = detail["documents"][0]["candidates"]
    assert all(candidate["preparation"]["state"] == "READY" for candidate in candidates)
    assert all(candidate["preparation"]["ready"] is True for candidate in candidates)

    original_answer_for = main.answer_for
    main.answer_for = lambda *_args, **_kwargs: {
        "answer": "관련 문서를 찾지 못했습니다.",
        "citations": [],
        "relevance": 0.0,
        "grounded": False,
        "retrieval_metadata": {},
    }
    try:
        comparison = client.post(
            f"/api/v1/rag-instances/{instance_id}/tuning/compare",
            json={"document_ids": [document_id], "question": "화성 이주 비용은 얼마인가요?"},
        ).json()
    finally:
        main.answer_for = original_answer_for
    unavailable = next(result for result in comparison["results"] if result["candidate_state"] == "NO_EVIDENCE")
    assert unavailable["citations"] == []
    assert "근거를 찾지 못" in unavailable["candidate_state_detail"]

    vote = client.post(
        f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote",
        json={"candidate_ids": [unavailable["candidate"]["id"]]},
    )
    assert vote.status_code == 422
    assert vote.json()["detail"]["code"] == "CANDIDATE_NOT_SELECTABLE"


def test_partial_candidate_index_failure_is_persisted_without_hiding_ready_candidates() -> None:
    instance_id = create_instance()
    original_prepare = main.prepare_candidate_index
    failed_once = False

    def prepare_with_one_failure(instance: main.RagInstance, candidate: main.Candidate) -> None:
        nonlocal failed_once
        if not failed_once:
            failed_once = True
            raise RuntimeError("test embedding endpoint timeout")
        original_prepare(instance, candidate)

    main.prepare_candidate_index = prepare_with_one_failure
    try:
        upload = client.post(
            f"/api/v1/rag-instances/{instance_id}/documents",
            json={"documents": [{"filename": "travel.txt", "content": "국내 출장 숙박비는 1박 10만원입니다."}]},
        ).json()
        assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    finally:
        main.prepare_candidate_index = original_prepare

    document_id = upload["documents"][0]["id"]
    detail = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    candidates = detail["documents"][0]["candidates"]
    failed = next(candidate for candidate in candidates if candidate["preparation"]["state"] == "FAILED")
    assert failed["preparation"]["error"] == "test embedding endpoint timeout"
    assert any(candidate["preparation"]["state"] == "READY" for candidate in candidates)

    comparison = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    ).json()
    failed_result = next(
        result for result in comparison["results"] if result["candidate"]["id"] == failed["id"]
    )
    assert failed_result["candidate_state"] == "FAILED"
    assert failed_result["citations"] == []

    job_id = upload["job"]["id"]
    assert client.get(f"/api/v1/rag-jobs/{job_id}").json()["can_retry"] is True
    retried = client.post(f"/api/v1/rag-jobs/{job_id}/retry")
    assert retried.status_code == 200
    assert wait_for_job(job_id)["state"] == "SUCCEEDED"


def test_tied_candidates_cannot_finalize() -> None:
    instance_id = create_instance()
    upload = client.post(f"/api/v1/rag-instances/{instance_id}/documents", json={"documents": [{"filename": "notes.txt", "content": "휴가는 연차 규정에 따라 신청합니다."}]})
    document_id = upload.json()["documents"][0]["id"]
    assert wait_for_job(upload.json()["job"]["id"])["state"] == "SUCCEEDED"
    comparison = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/compare", json={"document_ids": [document_id], "question": "휴가는 어떻게 신청해?"}).json()
    candidate_ids = [result["candidate"]["id"] for result in comparison["results"][:2]]
    client.post(f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote", json={"candidate_ids": candidate_ids})
    finalized = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/finalize", json={"document_id": document_id})
    assert finalized.status_code == 409
    assert finalized.json()["detail"]["code"] == "TUNING_TIED_OR_UNVOTED"


def test_ungrounded_query_is_refused() -> None:
    instance_id = create_instance()
    upload = client.post(f"/api/v1/rag-instances/{instance_id}/documents", json={"documents": [{"filename": "policy.txt", "content": "연차 휴가는 사전 승인이 필요합니다."}]})
    document_id = upload.json()["documents"][0]["id"]
    assert wait_for_job(upload.json()["job"]["id"])["state"] == "SUCCEEDED"
    comparison = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/compare", json={"document_ids": [document_id], "question": "연차 휴가"}).json()
    candidate_id = comparison["results"][0]["candidate"]["id"]
    client.post(f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote", json={"candidate_ids": [candidate_id]})
    client.post(f"/api/v1/rag-instances/{instance_id}/tuning/finalize", json={"document_id": document_id})
    response = client.post(f"/api/v1/rag-instances/{instance_id}/search", json={"document_ids": [document_id], "question": "화성 이주 비용은?", "sensitivity": "strict"})
    assert response.status_code == 200
    assert response.json()["grounded"] is False


def test_candidate_pipelines_use_distinct_persisted_chunks_and_citations() -> None:
    instance_id = create_instance()
    content = """제1조 목적
이 문서는 출장 절차와 비용 기준을 설명합니다. 신청은 사전에 승인받아야 합니다.

제2조 국내 출장
국내 출장 숙박비는 1박 10만원을 한도로 합니다. 교통비는 영수증 기준으로 정산합니다.

제3조 해외 출장
해외 출장 식비는 국가 등급에 따라 1일 80~150달러입니다. 항공료는 사전 승인을 받아야 합니다."""
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "travel-policy.txt", "content": content}]},
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    document_id = upload["documents"][0]["id"]
    document = main.INSTANCES[instance_id].documents[document_id]
    candidates = [main.CANDIDATES[candidate_id] for candidate_id in document.candidate_ids]
    semantic_candidate = next(candidate for candidate in candidates if candidate.chunking_strategy == "semantic")
    hierarchical_candidate = next(candidate for candidate in candidates if candidate.chunking_strategy == "hierarchical")
    assert len(semantic_candidate.segments) == 1
    assert len(hierarchical_candidate.segments) == 3
    assert {segment.id for segment in semantic_candidate.segments}.isdisjoint(
        {segment.id for segment in hierarchical_candidate.segments}
    )

    comparison = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    ).json()
    result = next(item for item in comparison["results"] if item["candidate"]["id"] == hierarchical_candidate.id)
    assert result["candidate"]["chunk_count"] == 3
    citation = result["citations"][0]
    assert citation["segment_id"] in {segment.id for segment in hierarchical_candidate.segments}
    assert client.get(citation["navigate_url"]).status_code == 200


def test_chunking_candidates_adapt_parameters_to_document_shape_and_expose_reasons() -> None:
    instance_id = create_instance()
    long_section = "세부 규정은 사전 승인과 증빙 보관을 요구합니다. " * 120
    structured_content = f"""제1조 목적
{long_section}

제2조 국내 출장
국내 출장 숙박비는 1박 10만원을 한도로 합니다.

제3조 해외 출장
해외 출장 식비는 국가 등급에 따라 달라집니다."""
    table_content = "항목 | 한도\n숙박비 | 10만원\n식비 | 150달러\n교통비 | 실비\n항공료 | 사전 승인"
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={
            "documents": [
                {"filename": "travel-policy.txt", "content": structured_content},
                {"filename": "travel-limits.csv", "content": table_content},
            ]
        },
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"

    detail = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    structured = next(document for document in detail["documents"] if document["filename"] == "travel-policy.txt")
    table = next(document for document in detail["documents"] if document["filename"] == "travel-limits.csv")
    assert structured["profile"] == "structured"
    assert structured["chunking_analysis"]["heading_count"] == 3
    assert table["profile"] == "table"

    hierarchical = next(candidate for candidate in structured["candidates"] if candidate["chunking_strategy"] == "hierarchical")
    semantic = next(candidate for candidate in structured["candidates"] if candidate["chunking_strategy"] == "semantic")
    table_candidate = next(candidate for candidate in table["candidates"] if candidate["chunking_strategy"] == "table")
    assert hierarchical["chunking_parameters"]["preserve_heading"] is True
    assert hierarchical["chunking_parameters"]["max_section_chars"] >= 800
    assert hierarchical["chunk_count"] > 3  # The oversized first section is split at the calculated limit.
    assert semantic["chunking_parameters"]["target_chars"] >= 420
    assert table_candidate["chunking_parameters"]["rows_per_chunk"] >= 3
    assert "표 형태" in table_candidate["selection_reason"]

    persisted = next(item for item in main.state_payload()["candidates"] if item["id"] == hierarchical["id"])
    assert persisted["chunking_parameters"] == hierarchical["chunking_parameters"]


def finalize_single_document(instance_id: str, content: str = "제7조 국내 출장 숙박비는 1박 10만원을 한도로 합니다.") -> str:
    upload = client.post(f"/api/v1/rag-instances/{instance_id}/documents", json={"documents": [{"filename": "policy.txt", "content": content}]})
    document_id = upload.json()["documents"][0]["id"]
    assert wait_for_job(upload.json()["job"]["id"])["state"] == "SUCCEEDED"
    comparison = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/compare", json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"}).json()
    winner_id = comparison["results"][0]["candidate"]["id"]
    assert client.post(f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote", json={"candidate_ids": [winner_id]}).status_code == 200
    assert client.post(f"/api/v1/rag-instances/{instance_id}/tuning/finalize", json={"document_id": document_id}).status_code == 200
    return document_id


def finalize_document_with_retrieval_config(instance_id: str, document_id: str, question: str, retrieval_config: str) -> None:
    comparison = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": question},
    ).json()
    winner = next(
        result for result in comparison["results"]
        if result["candidate_state"] == "READY" and result["candidate"]["retrieval_config"] == retrieval_config
    )
    assert client.post(
        f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote",
        json={"candidate_ids": [winner["candidate"]["id"]]},
    ).status_code == 200
    assert client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/finalize",
        json={"document_id": document_id},
    ).status_code == 200


def test_job_metadata_and_reuse_or_retune_contract() -> None:
    instance_id = create_instance()
    initial = client.post(f"/api/v1/rag-instances/{instance_id}/documents", json={"documents": [{"filename": "travel.txt", "content": "제7조 국내 출장 숙박비는 1박 10만원을 한도로 합니다."}], "pipeline_mode": "retune"})
    assert initial.status_code == 202
    job = initial.json()["job"]
    assert job["state"] in {"QUEUED", "PARSING", "GENERATING_CANDIDATES", "INDEXING"}
    completed_job = wait_for_job(job["id"])
    assert [stage["key"] for stage in completed_job["stages"]] == ["PARSING", "CANDIDATES", "INDEXING"]
    assert all(stage["state"] == "SUCCEEDED" for stage in completed_job["stages"])
    assert initial.json()["artifact"]["type"] == "PROCESSING_RUN"
    document_id = initial.json()["documents"][0]["id"]

    comparison = client.post(f"/api/v1/rag-instances/{instance_id}/tuning/compare", json={"document_ids": [document_id], "question": "숙박비는?"}).json()
    winner_id = comparison["results"][0]["candidate"]["id"]
    client.post(f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote", json={"candidate_ids": [winner_id]})
    client.post(f"/api/v1/rag-instances/{instance_id}/tuning/finalize", json={"document_id": document_id})

    options = client.get(f"/api/v1/rag-instances/{instance_id}/document-add-options")
    assert options.status_code == 200
    assert options.json()["default_mode"] == "reuse"
    assert options.json()["reusable_sources"][0]["document_id"] == document_id

    reused = client.post(f"/api/v1/rag-instances/{instance_id}/documents", json={"documents": [{"filename": "travel-extra.txt", "content": "출장은 사전 승인 후 진행합니다."}], "pipeline_mode": "reuse", "reuse_from_document_id": document_id})
    assert reused.status_code == 202
    assert wait_for_job(reused.json()["job"]["id"])["state"] == "SUCCEEDED"
    assert reused.json()["decision"]["next_action"] == "SEARCH_READY"
    refreshed = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    assert any(document["finalized_candidate_id"] for document in refreshed["documents"] if document["id"] != document_id)
    assert client.get(f"/api/v1/rag-instances/{instance_id}/jobs").json()["total"] == 2
    artifact_types = {item["type"] for item in client.get(f"/api/v1/rag-instances/{instance_id}/artifacts").json()["items"]}
    assert {"PROCESSING_RUN", "TUNING_COMPARISON", "PIPELINE_DECISION"} <= artifact_types


def test_citation_navigation_feedback_signal_retune_and_document_delete() -> None:
    instance_id = create_instance()
    document_id = finalize_single_document(instance_id)
    searched = client.post(f"/api/v1/rag-instances/{instance_id}/search", json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"})
    assert searched.status_code == 200
    search_body = searched.json()
    citation = search_body["citations"][0]
    viewer = client.get(citation["navigate_url"])
    assert viewer.status_code == 200
    assert viewer.json()["viewer"]["highlight"]["segment_id"] == citation["segment_id"]

    for _ in range(3):
        feedback = client.post(f"/api/v1/rag-instances/{instance_id}/feedback", json={"rating": -1, "artifact_id": search_body["artifact"]["id"], "document_ids": [document_id], "citation_ids": [citation["id"]]})
        assert feedback.status_code == 201
    assert feedback.json()["retuning_signal"]["recommended"] is True
    assert client.get(f"/api/v1/rag-instances/{instance_id}/feedback-summary").json()["negative_count"] == 3

    retune = client.post(f"/api/v1/rag-instances/{instance_id}/retune", json={"document_ids": [document_id], "reason": "부정 피드백 확인"})
    assert retune.status_code == 202
    assert retune.json()["next_action"] == "TUNE_DOCUMENT"
    assert wait_for_job(retune.json()["job"]["id"])["state"] == "SUCCEEDED"
    refreshed = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    assert len(refreshed["documents"][0]["candidates"]) == 9

    deleted = client.delete(f"/api/v1/rag-instances/{instance_id}/documents/{document_id}")
    assert deleted.status_code == 200
    assert deleted.json()["artifacts_with_unavailable_context"]


def test_job_and_comparison_are_restored_after_runtime_reload() -> None:
    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={
            "documents": [
                {
                    "filename": "restore.txt",
                    "content": "제7조 국내 출장 숙박비는 1박 10만원을 한도로 합니다.",
                }
            ]
        },
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    document_id = upload["documents"][0]["id"]
    comparison = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    ).json()
    winner_id = comparison["results"][0]["candidate"]["id"]
    assert client.post(
        f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote",
        json={"candidate_ids": [winner_id]},
    ).status_code == 200

    main.persist_state()
    main.INSTANCES.clear()
    main.CANDIDATES.clear()
    main.JOBS.clear()
    main.ROUNDS.clear()
    main.ARTIFACTS.clear()
    main.FEEDBACK.clear()
    main.restore_state()

    restored = client.get(f"/api/v1/rag-instances/{instance_id}")
    assert restored.status_code == 200
    body = restored.json()
    assert body["latest_job"]["state"] == "SUCCEEDED"
    assert body["latest_round"]["id"] == comparison["round"]["id"]
    assert body["latest_round"]["results"][0]["candidate"]["selection_count"] == 1
    assert main.CANDIDATES[winner_id].segments


def test_large_document_uses_bounded_sample_then_persists_full_reindex(monkeypatch) -> None:
    monkeypatch.setenv("RAG_PORTAL_COMPARISON_CHUNK_THRESHOLD", "3")
    assert client.get("/api/v1/large-document-policy").json()["comparison_chunk_threshold"] == 3
    instance_id = create_instance()
    content = "\n\n".join(
        f"제{index}조 국내 출장 숙박비는 1박 10만원을 한도로 하며 사전 승인이 필요합니다."
        for index in range(1, 11)
    )
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "large-policy.txt", "content": content}]},
    )
    assert upload.status_code == 202
    decision = upload.json()["decision"]
    assert decision["comparison_scope"] == "SAMPLE"
    assert decision["estimated_chunk_count"] > 3
    assert decision["selected_chunk_count"] == 3
    initial_job = wait_for_job(upload.json()["job"]["id"])
    document_id = upload.json()["documents"][0]["id"]
    assert initial_job["comparison_plans"][document_id]["scope"] == "SAMPLE"

    detail = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    document = detail["documents"][0]
    assert document["comparison"]["full_source_retained"] is True
    assert document["comparison"]["scope"] == "SAMPLE"
    assert document["comparison"]["selected_chunk_count"] == 3
    assert all(candidate["chunk_count"] <= 3 for candidate in document["candidates"])
    assert all(candidate["comparison"]["scope"] == "SAMPLE" for candidate in document["candidates"])

    comparison = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    ).json()
    selectable = next(result for result in comparison["results"] if result["candidate_state"] == "READY")
    assert client.post(
        f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote",
        json={"candidate_ids": [selectable["candidate"]["id"]]},
    ).status_code == 200
    finalized = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/finalize",
        json={"document_id": document_id},
    )
    assert finalized.status_code == 200
    reindex = finalized.json()["full_reindex"]
    assert reindex["required"] is True
    assert reindex["job"]["kind"] == "FULL_REINDEX"
    assert reindex["artifact"]["type"] == "FULL_REINDEX"
    full_job = wait_for_job(reindex["job"]["id"])
    assert full_job["kind"] == "FULL_REINDEX"
    assert [stage["key"] for stage in full_job["stages"]] == ["FULL_CHUNKING", "FULL_INDEXING"]
    assert full_job["state"] == "SUCCEEDED"

    searched = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    )
    assert searched.status_code == 200

    restored_detail = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    restored_document = restored_detail["documents"][0]
    assert restored_document["full_reindex"]["state"] == "SUCCEEDED"
    winner = restored_document["candidates"][0]
    assert winner["chunk_count"] == winner["comparison"]["estimated_chunk_count"]

    main.persist_state()
    main.INSTANCES.clear()
    main.CANDIDATES.clear()
    main.JOBS.clear()
    main.ROUNDS.clear()
    main.ARTIFACTS.clear()
    main.restore_state()
    snapshot_document = client.get(f"/api/v1/rag-instances/{instance_id}").json()["documents"][0]
    assert snapshot_document["comparison"]["scope"] == "SAMPLE"
    assert snapshot_document["full_reindex"]["job_id"] == reindex["job"]["id"]
    assert main.JOBS[reindex["job"]["id"]].kind == main.JobKind.FULL_REINDEX


def test_model_generation_uses_only_retrieved_context_and_persists_grounding_metadata(monkeypatch) -> None:
    received_contexts: list[dict] = []

    def fake_generator(*, question: str, contexts: list[dict]) -> GenerationResult:
        assert question == "국내 출장 숙박비 한도는?"
        received_contexts.extend(contexts)
        assert all(set(context) == {"segment_id", "text"} for context in contexts)
        return GenerationResult(
            sentences=[GenerationSentence(text="국내 출장 숙박비는 1박 10만원입니다.", citation_ids=[contexts[0]["segment_id"]])],
            provider="test-local-generator",
            model="test-grounded-model",
            latency_ms=7,
        )

    monkeypatch.setattr(main, "generate_grounded", fake_generator)
    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={"documents": [{"filename": "policy.txt", "content": "국내 출장 숙박비는 1박 10만원을 한도로 합니다."}]},
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    document_id = upload["documents"][0]["id"]
    comparison = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    ).json()
    comparable = next(result for result in comparison["results"] if result["candidate_state"] == "READY")
    assert comparable["generation"]["mode"] == "MODEL"
    assert comparable["generation"]["grounding_valid"] is True
    assert comparable["generation"]["provider"] == "test-local-generator"
    assert set(comparable["generation"]["sentence_citation_ids"][0]) <= set(comparable["generation"]["supplied_segment_ids"])
    assert client.post(
        f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote",
        json={"candidate_ids": [comparable["candidate"]["id"]]},
    ).status_code == 200
    assert client.post(f"/api/v1/rag-instances/{instance_id}/tuning/finalize", json={"document_id": document_id}).status_code == 200

    searched = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    ).json()
    assert searched["generation"]["mode"] == "MODEL"
    assert searched["generation"]["fallback"] is False
    assert searched["citations"][0]["segment_id"] in searched["generation"]["supplied_segment_ids"]
    artifact = client.get(f"/api/v1/rag-instances/{instance_id}/artifacts/{searched['artifact']['id']}").json()
    assert artifact["metadata"]["generation"]["provider"] == "test-local-generator"
    assert received_contexts


def test_invalid_or_failed_generation_streams_explicit_extractive_fallback(monkeypatch) -> None:
    def invalid_generator(*, question: str, contexts: list[dict]) -> GenerationResult:
        return GenerationResult(
            sentences=[
                GenerationSentence(text="근거 없는 문장입니다.", citation_ids=["not-a-supplied-segment"]),
            ],
            provider="test-local-generator",
            model="test-grounded-model",
            latency_ms=1,
        )

    monkeypatch.setattr(main, "generate_grounded", invalid_generator)
    instance_id = create_instance()
    document_id = finalize_single_document(instance_id)
    response = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["generation"]["mode"] == "EXTRACTIVE_FALLBACK"
    assert body["generation"]["fallback_reason"] == "INVALID_GROUNDING"
    assert "근거 없는 문장" not in body["answer"]
    assert body["citations"]

    def unavailable_generator(*, question: str, contexts: list[dict]) -> GenerationResult:
        raise GenerationEndpointError("generator offline")

    monkeypatch.setattr(main, "generate_grounded", unavailable_generator)
    stream = client.get(
        f"/api/v1/rag-instances/{instance_id}/search/stream",
        params=[("question", "국내 출장 숙박비 한도는?"), ("document_ids", document_id)],
    )
    assert stream.status_code == 200
    assert 'event: status' in stream.text
    assert '"phase": "RETRIEVING"' in stream.text
    assert '"phase": "ANSWER_READY"' in stream.text
    assert "event: citations" in stream.text
    assert "event: token" in stream.text
    assert '"delivery": "BUFFERED_REPLAY"' in stream.text
    assert '"mode": "EXTRACTIVE_FALLBACK"' in stream.text
    assert '"fallback_reason": "GENERATOR_UNAVAILABLE"' in stream.text
    assert '"server_cancellation": "disconnect_stops_future_events_but_does_not_interrupt_active_provider_request"' in stream.text


def test_empty_generation_is_rejected_and_uses_cited_fallback(monkeypatch) -> None:
    def empty_generator(*, question: str, contexts: list[dict]) -> GenerationResult:
        return GenerationResult(sentences=[], provider="test-local-generator", model="test-grounded-model", latency_ms=1)

    monkeypatch.setattr(main, "generate_grounded", empty_generator)
    instance_id = create_instance()
    document_id = finalize_single_document(instance_id)
    response = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["generation"]["mode"] == "EXTRACTIVE_FALLBACK"
    assert body["generation"]["fallback_reason"] == "INVALID_GROUNDING"
    assert body["citations"]
    assert set(item["segment_id"] for item in body["citations"]) <= set(body["generation"]["supplied_segment_ids"])

    stream = client.get(
        f"/api/v1/rag-instances/{instance_id}/search/stream",
        params=[("question", "국내 출장 숙박비 한도는?"), ("document_ids", document_id)],
    )
    assert stream.status_code == 200
    assert '"fallback_reason": "INVALID_GROUNDING"' in stream.text


def test_full_reindex_search_policy_is_explicit_and_shared_by_rest_and_sse() -> None:
    instance_id = create_instance()
    document_id = finalize_single_document(instance_id)
    document = main.INSTANCES[instance_id].documents[document_id]
    document.comparison_scope = main.ComparisonScope.SAMPLE
    document.full_reindex_job_id = "full-reindex-job"
    document.full_reindex_state = main.JobState.QUEUED

    detail = client.get(f"/api/v1/rag-instances/{instance_id}").json()["documents"][0]["full_reindex"]
    assert detail["search_policy"] == "BLOCK_UNTIL_FULL_INDEX_SUCCEEDED"
    assert detail["search_eligible"] is False
    assert detail["next_action"] == "WAIT_FOR_FULL_REINDEX"

    request = {"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"}
    rest = client.post(f"/api/v1/rag-instances/{instance_id}/search", json=request)
    assert rest.status_code == 409
    assert rest.json()["detail"]["code"] == "FULL_REINDEX_PENDING"
    assert rest.json()["detail"]["details"] == detail
    stream = client.get(
        f"/api/v1/rag-instances/{instance_id}/search/stream",
        params=[("question", request["question"]), ("document_ids", document_id)],
    )
    assert stream.status_code == 409
    assert stream.json()["detail"] == rest.json()["detail"]

    document.full_reindex_state = main.JobState.FAILED
    failed = client.post(f"/api/v1/rag-instances/{instance_id}/search", json=request)
    assert failed.status_code == 409
    assert failed.json()["detail"]["code"] == "FULL_REINDEX_FAILED"
    assert failed.json()["detail"]["details"]["next_action"] == "RETRY_FULL_REINDEX"


def test_multi_document_search_normalizes_per_document_scores_and_groups_grounded_citations(monkeypatch) -> None:
    received_contexts: list[dict] = []

    def fake_generator(*, question: str, contexts: list[dict]) -> GenerationResult:
        received_contexts.extend(contexts)
        sentences = [GenerationSentence(text="첫 번째 근거입니다.", citation_ids=[contexts[0]["segment_id"]])]
        if len(contexts) > 1:
            sentences.append(GenerationSentence(text="두 번째 근거입니다.", citation_ids=[contexts[1]["segment_id"]]))
        return GenerationResult(
            sentences=sentences,
            provider="test-local-generator",
            model="test-grounded-model",
            latency_ms=3,
        )

    monkeypatch.setattr(main, "generate_grounded", fake_generator)
    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={
            "documents": [
                {"filename": "approval.txt", "content": "출장은 사전에 승인을 받아야 합니다."},
                {"filename": "meal.txt", "content": "해외 출장 식비는 하루 150달러를 한도로 합니다."},
            ]
        },
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    approval_id, meal_id = (document["id"] for document in upload["documents"])
    question = "출장 사전 승인 요건과 해외 식비 한도는?"
    finalize_document_with_retrieval_config(instance_id, approval_id, question, "dense")
    finalize_document_with_retrieval_config(instance_id, meal_id, question, "hybrid_rerank")
    received_contexts.clear()

    response = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [approval_id, meal_id], "question": question},
    )
    assert response.status_code == 200
    body = response.json()
    metadata = body["retrieval_metadata"]
    assert body["retrieval_config"] == "per_document"
    assert metadata["mode"] == "MULTI_DOCUMENT"
    assert metadata["merge"]["strategy"] == "per_document_max_normalized_global_merge"
    assert [item["retrieval_config"] for item in metadata["documents"]] == ["dense", "hybrid_rerank"]
    assert metadata["documents"][1]["rerank"]["applied"] is True
    assert all(
        0 <= candidate["normalized_score"] <= 1
        for document in metadata["documents"]
        for candidate in document["top_candidates"]
    )
    global_scores = [candidate["normalized_score"] for candidate in metadata["global_candidates"]]
    assert global_scores == sorted(global_scores, reverse=True)
    assert [citation["document_id"] for citation in body["citations"]] == [
        metadata["global_candidates"][0]["document_id"],
        metadata["global_candidates"][1]["document_id"],
    ]
    assert {group["document_id"] for group in body["grouped_citations"]} == {approval_id, meal_id}
    assert len(received_contexts) == 2
    artifact = client.get(f"/api/v1/rag-instances/{instance_id}/artifacts/{body['artifact']['id']}").json()
    assert artifact["metadata"]["retrieval_metadata"]["mode"] == "MULTI_DOCUMENT"
    assert artifact["payload"]["grouped_citations"] == body["grouped_citations"]
    snapshot = main.STATE_STORE.load()
    snapshot_artifact = next(item for item in snapshot["artifacts"] if item["id"] == body["artifact"]["id"])
    assert snapshot_artifact["payload"]["retrieval_metadata"]["mode"] == "MULTI_DOCUMENT"


def test_multi_document_invalid_generation_falls_back_only_to_selected_evidence(monkeypatch) -> None:
    def invalid_generator(*, question: str, contexts: list[dict]) -> GenerationResult:
        return GenerationResult(
            sentences=[GenerationSentence(text="검증되지 않은 모델 문장입니다.", citation_ids=["missing-segment"])],
            provider="test-local-generator",
            model="test-grounded-model",
            latency_ms=1,
        )

    monkeypatch.setattr(main, "generate_grounded", invalid_generator)
    instance_id = create_instance()
    upload = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={
            "documents": [
                {"filename": "approval.txt", "content": "출장은 사전에 승인을 받아야 합니다."},
                {"filename": "meal.txt", "content": "해외 출장 식비는 하루 150달러를 한도로 합니다."},
            ]
        },
    ).json()
    assert wait_for_job(upload["job"]["id"])["state"] == "SUCCEEDED"
    approval_id, meal_id = (document["id"] for document in upload["documents"])
    question = "출장 사전 승인 요건과 해외 식비 한도는?"
    finalize_document_with_retrieval_config(instance_id, approval_id, question, "dense")
    finalize_document_with_retrieval_config(instance_id, meal_id, question, "hybrid_rerank")

    response = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [approval_id, meal_id], "question": question},
    )
    assert response.status_code == 200
    body = response.json()
    assert body["generation"]["mode"] == "EXTRACTIVE_FALLBACK"
    assert body["generation"]["fallback_reason"] == "INVALID_GROUNDING"
    assert "검증되지 않은 모델 문장" not in body["answer"]
    assert len(body["citations"]) == 2
    assert set(citation["segment_id"] for citation in body["citations"]) <= set(body["generation"]["supplied_segment_ids"])


def test_retuning_recommendation_is_explainable_and_persists_baseline_and_outcome() -> None:
    instance_id = create_instance()
    document_id = finalize_single_document(instance_id)
    searched = client.post(
        f"/api/v1/rag-instances/{instance_id}/search",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    ).json()
    # A fallback/partial benchmark is visible as runtime context only. It must
    # never become a fabricated model-quality score in the recommendation.
    main.BENCHMARK_RUNS.append(
        {
            "run": {"id": "fallback-benchmark", "created_at": main.now(), "corpus_label": "test"},
            "results": [{"model_id": "BGE-M3", "status": "FALLBACK", "provider": "local-hash-fallback", "dimension": 96}],
        }
    )
    for _ in range(2):
        response = client.post(
            f"/api/v1/rag-instances/{instance_id}/feedback",
            json={
                "rating": -1,
                "artifact_id": searched["artifact"]["id"],
                "document_ids": [document_id],
                "citation_ids": [citation["id"] for citation in searched["citations"]],
            },
        )
        assert response.status_code == 201

    recommendation = client.get(f"/api/v1/rag-instances/{instance_id}/retuning-recommendation")
    assert recommendation.status_code == 200
    signal = recommendation.json()
    assert signal["version"] == main.RETUNING_SIGNAL_VERSION
    assert signal["recommended"] is True
    assert signal["inputs"]["feedback"]["negative_recency_weight"] == 2.0
    assert signal["inputs"]["benchmark_provider"]["status"] == "NOT_RELEASE_EVIDENCE"
    assert signal["inputs"]["benchmark_provider"]["used_for_recommendation_score"] is False
    assert "BENCHMARK_NOT_QUALITY_EVIDENCE" in signal["recommendation_reasons"]
    assert signal["baseline_snapshot"]["document_ids"] == [document_id]
    assert signal["baseline_snapshot"]["selected_pipelines"][0]["pipeline"]["finalized"] is True
    summary = client.get(f"/api/v1/rag-instances/{instance_id}/feedback-summary").json()
    assert summary["retuning_recommendation"]["version"] == signal["version"]

    retune = client.post(
        f"/api/v1/rag-instances/{instance_id}/retune",
        json={"document_ids": [document_id], "reason": "feedback and evidence review"},
    )
    assert retune.status_code == 202
    retune_body = retune.json()
    baseline = client.get(
        f"/api/v1/rag-instances/{instance_id}/artifacts/{retune_body['baseline_artifact']['id']}"
    ).json()
    assert baseline["type"] == "RETUNING_BASELINE"
    assert baseline["payload"]["recommendation"]["version"] == main.RETUNING_SIGNAL_VERSION
    assert baseline["payload"]["selected_pipelines"][0]["document_id"] == document_id
    assert wait_for_job(retune_body["job"]["id"])["state"] == "SUCCEEDED"

    compared = client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/compare",
        json={"document_ids": [document_id], "question": "국내 출장 숙박비 한도는?"},
    )
    assert compared.status_code == 200
    comparison = compared.json()
    assert comparison["retuning_outcome_artifact"] is not None
    outcome = client.get(
        f"/api/v1/rag-instances/{instance_id}/artifacts/{comparison['retuning_outcome_artifact']['id']}"
    ).json()
    assert outcome["type"] == "RETUNING_OUTCOME"
    assert outcome["metadata"]["baseline_artifact_id"] == retune_body["baseline_artifact"]["id"]
    assert outcome["metadata"]["selection_state"] == "PENDING_USER_VOTE"
    assert outcome["payload"]["comparison_observations"]["candidate_count"] == len(comparison["results"])

    winner = next(result for result in comparison["results"] if result["candidate_state"] == "READY")
    assert client.post(
        f"/api/v1/tuning-rounds/{comparison['round']['id']}/vote",
        json={"candidate_ids": [winner["candidate"]["id"]]},
    ).status_code == 200
    assert client.post(
        f"/api/v1/rag-instances/{instance_id}/tuning/finalize",
        json={"document_id": document_id},
    ).status_code == 200
    finalized_outcome = client.get(
        f"/api/v1/rag-instances/{instance_id}/artifacts/{outcome['id']}"
    ).json()
    assert finalized_outcome["metadata"]["selection_state"] == "FINALIZED"
    assert document_id in finalized_outcome["metadata"]["selected_pipelines"]
    snapshot = main.state_payload()
    assert {"RETUNING_BASELINE", "RETUNING_OUTCOME"} <= {artifact["type"] for artifact in snapshot["artifacts"]}


def test_source_provenance_dedup_and_reparse_retry_preserve_artifact_history(monkeypatch, tmp_path) -> None:
    monkeypatch.setenv("RAG_SOURCE_STORAGE_PATH", str(tmp_path / "sources"))
    instance_id = create_instance()
    uploaded = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents",
        json={
            "documents": [
                {"filename": "policy-a.txt", "content": "국내 출장 숙박비는 1박 10만원을 한도로 합니다."},
                {"filename": "policy-b.txt", "content": "국내 출장 숙박비는 1박 10만원을 한도로 합니다."},
            ]
        },
    )
    assert uploaded.status_code == 202
    upload_body = uploaded.json()
    assert wait_for_job(upload_body["job"]["id"])["state"] == "SUCCEEDED"
    first_id, second_id = (document["id"] for document in upload_body["documents"])
    detail = client.get(f"/api/v1/rag-instances/{instance_id}").json()
    first, second = detail["documents"]
    assert first["source"]["checksum_sha256"] == second["source"]["checksum_sha256"]
    assert first["source"]["storage_key"] == second["source"]["storage_key"]
    assert first["source"]["deduplication"]["status"] == "STORED"
    assert second["source"]["deduplication"]["status"] == "DEDUPLICATED"
    assert second["source"]["deduplication"]["deduplicated_from_document_id"] == first_id
    assert first["provenance"]["parser"]["version"] == main.PARSER_PIPELINE_VERSION
    assert first["provenance"]["chunking"]["version"] == main.CHUNKING_PIPELINE_VERSION
    assert first["provenance"]["model"]["version"] == main.MODEL_PROVENANCE_VERSION
    assert first["provenance"]["model"]["embedding_model"] == "BGE-M3"

    original_start = main.start_processing_job
    monkeypatch.setattr(main, "start_processing_job", lambda *args, **kwargs: None)
    reparse = client.post(
        f"/api/v1/rag-instances/{instance_id}/documents/{first_id}/reparse",
        json={"reason": "parser release verification"},
    )
    assert reparse.status_code == 202
    reparse_body = reparse.json()
    assert reparse_body["job"]["kind"] == "REPARSE"
    baseline = client.get(
        f"/api/v1/rag-instances/{instance_id}/artifacts/{reparse_body['baseline_artifact']['id']}"
    ).json()
    assert baseline["type"] == "REPARSE_BASELINE"
    assert baseline["payload"]["document"]["source"]["checksum_sha256"] == first["source"]["checksum_sha256"]
    cancelled = client.post(f"/api/v1/rag-jobs/{reparse_body['job']['id']}/cancel")
    assert cancelled.status_code == 200
    # The queue is intentionally paused above; let the existing worker boundary
    # observe the cancellation before exercising its normal retry endpoint.
    main.run_processing_job(reparse_body["job"]["id"], main.PipelineMode.RETUNE)
    assert client.get(f"/api/v1/rag-jobs/{reparse_body['job']['id']}").json()["state"] == "CANCELLED"
    monkeypatch.setattr(main, "start_processing_job", original_start)
    retried = client.post(f"/api/v1/rag-jobs/{reparse_body['job']['id']}/retry")
    assert retried.status_code == 200
    assert retried.json()["kind"] == "REPARSE"
    assert wait_for_job(reparse_body["job"]["id"])["state"] == "SUCCEEDED"

    reparsed = client.get(f"/api/v1/rag-instances/{instance_id}").json()["documents"]
    reparsed_first = next(document for document in reparsed if document["id"] == first_id)
    assert reparsed_first["source"]["storage_key"] == first["source"]["storage_key"]
    assert reparsed_first["provenance"]["parser"]["parse_revision"] == 2
    run_artifact = client.get(
        f"/api/v1/rag-instances/{instance_id}/artifacts/{reparse_body['artifact']['id']}"
    ).json()
    assert run_artifact["type"] == "REPARSE_RUN"
    assert run_artifact["metadata"]["baseline_artifact_id"] == reparse_body["baseline_artifact"]["id"]
    assert run_artifact["metadata"]["reparsed_provenance"][first_id]["parser"]["parse_revision"] == 2
    assert second_id in main.INSTANCES[instance_id].documents
    monkeypatch.setenv("RAG_SOURCE_STORAGE_BACKEND", "object")
    monkeypatch.setenv("RAG_OBJECT_STORAGE_ENDPOINT", "https://object-gateway.test")
    monkeypatch.setenv("RAG_OBJECT_STORAGE_BUCKET", "rag-sources")
    assert source_storage().backend == "http_object_storage"
